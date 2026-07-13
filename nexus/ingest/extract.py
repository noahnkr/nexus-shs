"""Text extraction.

Turn a raw source into text the classifier can read. Lossless originals: file-based ingest
archives the binary original (the pipeline does this); text-only ingest treats the note
body as authoritative and archives nothing.

Text formats are handled here with no extra dependencies. Binary formats dispatch to a
per-suffix extractor (all in-process, pure-Python — no external service): `.pdf` via
pdfminer.six, `.docx` via python-docx. The dispatch point is `extract_text`.

Google Docs are NOT handled here: they are not a byte format but a Drive resource. Export
them to text/markdown (or .docx) in the Drive connector's fetch step, so they arrive as
text or as a .docx the extractor below already reads.
"""

from __future__ import annotations

import re
from pathlib import Path

_TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".rst", ".csv", ".json", ".log"}
_TAG = re.compile(r"<[^>]+>")


def extract_text(source: Path) -> str:
    """Return plain text for `source`, dispatching on suffix."""
    suffix = source.suffix.lower()
    if suffix in _TEXT_SUFFIXES:
        return source.read_text(encoding="utf-8", errors="replace")
    if suffix in {".html", ".htm"}:
        return _TAG.sub(" ", source.read_text(encoding="utf-8", errors="replace"))
    if suffix == ".pdf":
        return _extract_pdf(source)
    if suffix == ".docx":
        return _extract_docx(source)
    raise NotImplementedError(
        f"no extractor registered for '{suffix}'. Add one here for binary formats; "
        "text formats are handled natively."
    )


def _extract_pdf(source: Path) -> str:
    """Extract the embedded text layer of a digital PDF via pdfminer.six.

    pdfminer reconstructs words from glyph positions, which survives the design-tool PDFs
    (Canva/Illustrator exports) whose text layers explode into per-character runs under
    naive extraction. Image-only (scanned) PDFs have no text layer and yield an empty
    result; we raise so a human notices rather than silently drafting an empty note. OCR
    (Tesseract) is a separate concern — add it here only if scanned docs are in scope.
    """
    from pdfminer.high_level import extract_text as pdf_extract_text

    text = (pdf_extract_text(str(source)) or "").replace("�", "").strip()
    if not text:
        raise ValueError(
            f"'{source.name}' yielded no extractable text (likely a scanned/image "
            "PDF). OCR is not wired in; extract text upstream or add an OCR extractor."
        )
    return _repair_exploded_text(text)


def _repair_exploded_text(text: str) -> str:
    """Collapse per-character letter spacing some PDF text layers emit.

    Certain PDFs encode each glyph as its own text run, so extraction yields
    'F r e q u e n t l y   A s k e d' — which tokenizes into single letters and makes the
    note body useless for search. Detected by the share of single-char tokens; repaired by
    treating runs of 2+ spaces as word breaks and removing the single spaces within words.
    """
    tokens = text.split()
    if not tokens or sum(len(t) == 1 for t in tokens) / len(tokens) < 0.7:
        return text
    lines = []
    for line in text.splitlines():
        words = re.split(r" {2,}", line)
        lines.append(" ".join(w.replace(" ", "") for w in words))
    return "\n".join(lines)


def _extract_docx(source: Path) -> str:
    """Extract paragraph and table text from a .docx via python-docx."""
    from docx import Document

    doc = Document(str(source))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()
